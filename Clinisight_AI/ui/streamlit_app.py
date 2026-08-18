import os
import uuid
from datetime import datetime
from io import BytesIO

import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF

DEFAULT_API_BASE = os.environ.get("CLINISIGHT_API_BASE", "http://localhost:8081")

# Mirrors the categories in functions/symptom_extractor.py::SYMPTOMS, with a
# distinct, colorful emoji per symptom (emoji render in full color, unlike
# the monochrome Material symbol font, so each pill/badge stays legible and
# visually distinct even inside a 90-item list).
SYMPTOM_ICONS_BY_CATEGORY = {
    "general": {
        "fever": "🤒", "chills": "🥶", "night sweats": "💦", "sweating": "😓",
        "fatigue": "😴", "malaise": "😔", "weakness": "😩",
        "dehydration": "🌵", "weight loss": "📉", "weight gain": "📈",
        "loss of appetite": "🍽️",
    },
    "neurological": {
        "headache": "🤕", "migraine": "🤯", "dizziness": "😵‍💫",
        "lightheadedness": "😵", "fainting": "💫", "confusion": "😕",
        "memory loss": "🧠", "numbness": "🫥", "tingling": "⚡",
        "seizures": "🌩️", "tremors": "🤲", "blurred vision": "👓",
        "double vision": "👀", "loss of taste": "👅", "loss of smell": "👃",
        "sensitivity to light": "💡", "sensitivity to sound": "🔊",
        "slurred speech": "🗣️", "difficulty speaking": "🤐",
    },
    "pain": {
        "chest pain": "🫀", "chest tightness": "🫁", "abdominal pain": "🫃",
        "stomach pain": "😖", "stomach ache": "😖", "abdominal cramps": "🌀",
        "back pain": "🦴", "neck pain": "💢", "joint pain": "🦵",
        "muscle pain": "💪", "muscle aches": "💪", "muscle weakness": "💪",
        "body aches": "🤕", "body ache": "🤕", "ear pain": "👂",
        "eye pain": "👁️", "throat pain": "🔥", "pelvic pain": "💢",
        "cramps": "🌀", "stiffness": "🪨", "pain": "💢",
    },
    "respiratory": {
        "shortness of breath": "🫁", "difficulty breathing": "😤",
        "wheezing": "🌬️", "dry cough": "😮‍💨", "cough": "😷",
        "sore throat": "🗣️", "runny nose": "🤧", "stuffy nose": "👃",
        "nasal congestion": "🤧", "sneezing": "🤧", "congestion": "😤",
    },
    "cardiovascular": {
        "palpitations": "💓", "irregular heartbeat": "💔",
        "high blood pressure": "📈", "low blood pressure": "📉",
    },
    "gastrointestinal": {
        "nausea": "🤢", "vomiting": "🤮", "diarrhea": "💩",
        "constipation": "🚽", "bloating": "🎈", "indigestion": "😣",
        "heartburn": "🔥", "blood in stool": "🩸",
    },
    "musculoskeletal_skin": {
        "swelling": "🫧", "rash": "🔴", "itching": "🪲", "hives": "🟠",
        "dry skin": "🏜️", "bruising": "🟣", "redness": "🔴",
    },
    "mental_health": {
        "anxiety": "😰", "depression": "😞", "insomnia": "🌙",
        "irritability": "😠", "mood swings": "🎭",
    },
}
DEFAULT_SYMPTOM_ICON = "🩺"
SYMPTOM_ICONS = {
    symptom: icon
    for category in SYMPTOM_ICONS_BY_CATEGORY.values()
    for symptom, icon in category.items()
}
KNOWN_SYMPTOMS = list(SYMPTOM_ICONS.keys())
SEVERITY_COLORS = {"Mild": "green", "Moderate": "orange", "Severe": "red"}

st.set_page_config(
    page_title="ClinInsight AI",
    page_icon=":material/health_metrics:",
    layout="wide",
)

st.session_state.setdefault("api_base", DEFAULT_API_BASE)
st.session_state.setdefault("history", [])
st.session_state.setdefault("current_id", None)


@st.cache_data(ttl=8, show_spinner=False)
def check_backend(api_base: str) -> tuple[bool, str]:
    try:
        response = requests.get(api_base, timeout=3)
        response.raise_for_status()
        return True, response.json().get("message", "Backend reachable")
    except Exception as exc:
        return False, str(exc)


def _post_diagnosis(api_base: str, description: str, force_refresh: bool) -> dict:
    response = requests.post(
        f"{api_base.rstrip('/')}/diagnosis",
        json={"description": description, "force_refresh": force_refresh},
        timeout=90,
    )
    response.raise_for_status()
    return response.json()


@st.cache_data(ttl="15m", max_entries=50, show_spinner=False)
def _call_diagnosis_api_cached(api_base: str, description: str) -> dict:
    return _post_diagnosis(api_base, description, force_refresh=False)


def call_diagnosis_api(api_base: str, description: str, force_refresh: bool = False) -> dict:
    if force_refresh:
        return _post_diagnosis(api_base, description, force_refresh=True)
    return _call_diagnosis_api_cached(api_base, description)


def clear_server_cache(api_base: str) -> dict:
    response = requests.post(f"{api_base.rstrip('/')}/cache/clear", timeout=15)
    response.raise_for_status()
    return response.json()


@st.cache_data(ttl=5, show_spinner=False)
def get_cache_stats(api_base: str) -> dict | None:
    try:
        response = requests.get(f"{api_base.rstrip('/')}/cache/stats", timeout=5)
        response.raise_for_status()
        return response.json()
    except Exception:
        return None


def build_report_markdown(entry: dict) -> str:
    result = entry["result"]
    lines = [
        "# ClinInsight AI — analysis report",
        f"*Generated {entry['timestamp']}*",
        "",
        "## Reported symptoms",
        entry["description"],
        "",
        "## Detected keywords",
        ", ".join(result.get("symptom", [])) or "None detected",
        "",
        "## AI diagnosis overview",
        result.get("diagnosis", ""),
        "",
        "## Literature summary",
        result.get("pubmed_summary", ""),
        "",
        "## Sources",
    ]
    for source in result.get("sources", []):
        authors = ", ".join(source.get("authors", []))
        lines.append(
            f"- [{source.get('title', 'Untitled')}]({source.get('article_url', '')}) "
            f"— {authors} ({source.get('publication_date', 'n.d.')})"
        )

    web_sources = result.get("web_sources", [])
    if web_sources:
        lines += ["", "## Web research"]
        for source in web_sources:
            lines.append(f"- [{source.get('title', 'Untitled')}]({source.get('url', '')})")

    lines += [
        "",
        "_This report is AI-generated and for informational purposes only. "
        "Consult a licensed clinician for medical advice._",
    ]
    return "\n".join(lines)


DISCLAIMER = (
    "This report is AI-generated and for informational purposes only. "
    "Consult a licensed clinician for medical advice."
)


def build_report_text(entry: dict) -> str:
    result = entry["result"]
    lines = [
        "CLININSIGHT AI - ANALYSIS REPORT",
        f"Generated {entry['timestamp']}",
        "",
        "REPORTED SYMPTOMS",
        "-----------------",
        entry["description"],
        "",
        "DETECTED KEYWORDS",
        "-----------------",
        ", ".join(result.get("symptom", [])) or "None detected",
        "",
        "AI DIAGNOSIS OVERVIEW",
        "----------------------",
        result.get("diagnosis", "") or "No diagnosis returned.",
        "",
        "LITERATURE SUMMARY",
        "-------------------",
        result.get("pubmed_summary", "") or "No summary available.",
        "",
        "SOURCES",
        "-------",
    ]
    sources = result.get("sources", [])
    if not sources:
        lines.append("No supporting literature was returned for this query.")
    for source in sources:
        authors = ", ".join(source.get("authors", []))
        lines.append(
            f"- {source.get('title', 'Untitled')} ({authors}, "
            f"{source.get('publication_date', 'n.d.')})"
        )
        if source.get("article_url"):
            lines.append(f"  {source['article_url']}")

    web_sources = result.get("web_sources", [])
    if web_sources:
        lines += ["", "WEB RESEARCH", "------------"]
        for source in web_sources:
            lines.append(f"- {source.get('title', 'Untitled')}")
            if source.get("url"):
                lines.append(f"  {source['url']}")

    lines += ["", DISCLAIMER]
    return "\n".join(lines)


def _add_docx_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_report_docx(entry: dict) -> bytes:
    result = entry["result"]
    document = Document()

    document.add_heading("ClinInsight AI — Analysis Report", level=0)
    caption = document.add_paragraph(f"Generated {entry['timestamp']}")
    caption.runs[0].italic = True

    document.add_heading("Reported symptoms", level=1)
    document.add_paragraph(entry["description"])

    document.add_heading("Detected keywords", level=1)
    document.add_paragraph(", ".join(result.get("symptom", [])) or "None detected")

    document.add_heading("AI diagnosis overview", level=1)
    document.add_paragraph(result.get("diagnosis", "") or "No diagnosis returned.")

    document.add_heading("Literature summary", level=1)
    document.add_paragraph(result.get("pubmed_summary", "") or "No summary available.")

    document.add_heading("Sources", level=1)
    sources = result.get("sources", [])
    if not sources:
        document.add_paragraph("No supporting literature was returned for this query.")
    for source in sources:
        authors = ", ".join(source.get("authors", []))
        bullet = document.add_paragraph(style="List Bullet")
        run = bullet.add_run(source.get("title", "Untitled"))
        run.bold = True
        bullet.add_run(f" — {authors} ({source.get('publication_date', 'n.d.')})")
        if source.get("article_url"):
            link_paragraph = document.add_paragraph()
            _add_docx_hyperlink(link_paragraph, source["article_url"], source["article_url"])

    web_sources = result.get("web_sources", [])
    if web_sources:
        document.add_heading("Web research", level=1)
        for source in web_sources:
            bullet = document.add_paragraph(style="List Bullet")
            bullet.add_run(source.get("title", "Untitled")).bold = True
            if source.get("url"):
                link_paragraph = document.add_paragraph()
                _add_docx_hyperlink(link_paragraph, source["url"], source["url"])

    footer = document.add_paragraph()
    footer_run = footer.add_run(DISCLAIMER)
    footer_run.italic = True

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_PDF_CHAR_MAP = {
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "–": "-",
    "—": "-",
    "…": "...",
    "•": "-",
    " ": " ",
}


def _pdf_safe(text: str) -> str:
    for src, dest in _PDF_CHAR_MAP.items():
        text = text.replace(src, dest)
    return text.encode("latin-1", "replace").decode("latin-1")


def build_report_pdf(entry: dict) -> bytes:
    result = entry["result"]
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def line(text: str, size: int = 10, style: str = "", h: int = 6, link: str | None = None) -> None:
        pdf.set_font("Helvetica", style, size)
        pdf.multi_cell(0, h, _pdf_safe(text), new_x="LMARGIN", new_y="NEXT", link=link)

    line("ClinInsight AI - Analysis Report", size=16, style="B", h=10)
    pdf.set_text_color(100, 100, 100)
    line(f"Generated {entry['timestamp']}", size=9, style="I")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    def section(title: str, body: str) -> None:
        line(title, size=12, style="B", h=8)
        line(body or "-", size=10)
        pdf.ln(2)

    section("Reported symptoms", entry["description"])
    section("Detected keywords", ", ".join(result.get("symptom", [])) or "None detected")
    section("AI diagnosis overview", result.get("diagnosis", ""))
    section("Literature summary", result.get("pubmed_summary", ""))

    line("Sources", size=12, style="B", h=8)
    sources = result.get("sources", [])
    if not sources:
        line("No supporting literature was returned for this query.", size=10)
    for source in sources:
        authors = ", ".join(source.get("authors", []))
        line(f"- {source.get('title', 'Untitled')}", size=10, style="B")
        line(f"{authors} ({source.get('publication_date', 'n.d.')})", size=9)
        if source.get("article_url"):
            pdf.set_text_color(30, 100, 200)
            line(source["article_url"], size=9, link=source["article_url"])
            pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

    web_sources = result.get("web_sources", [])
    if web_sources:
        pdf.ln(2)
        line("Web research", size=12, style="B", h=8)
        for source in web_sources:
            line(f"- {source.get('title', 'Untitled')}", size=10, style="B")
            line(source.get("content", "No summary available."), size=9)
            if source.get("url"):
                pdf.set_text_color(30, 100, 200)
                line(source["url"], size=9, link=source["url"])
                pdf.set_text_color(0, 0, 0)
            pdf.ln(1)

    pdf.ln(3)
    line(DISCLAIMER, size=8, style="I", h=5)

    return bytes(pdf.output())


def render_result(entry: dict) -> None:
    result = entry["result"]
    symptoms = result.get("symptom", [])
    sources = result.get("sources", [])
    web_sources = result.get("web_sources", [])

    with st.container(horizontal=True):
        st.caption(f"Analyzed {entry['timestamp']}")
        if entry.get("severity"):
            st.badge(entry["severity"], color=SEVERITY_COLORS.get(entry["severity"], "blue"))
        cache_hits = result.get("cache")
        if cache_hits:
            if all(cache_hits.values()):
                st.badge("Served from cache", icon=":material/bolt:", color="green")
            elif any(cache_hits.values()):
                st.badge("Partially cached", icon=":material/bolt:", color="orange")
            else:
                st.badge("Freshly generated", icon=":material/auto_awesome:", color="blue")

    diagnosis_tab, literature_tab, web_tab, summary_tab, report_tab = st.tabs(
        [
            "Diagnosis",
            "Literature review",
            "Web research",
            "Summary",
            "Report",
        ]
    )

    with diagnosis_tab:
        if symptoms:
            with st.container(horizontal=True):
                for symptom in symptoms:
                    icon = SYMPTOM_ICONS.get(symptom, DEFAULT_SYMPTOM_ICON)
                    st.badge(symptom, icon=icon, color="blue")
        else:
            st.caption("No recognized symptom keywords were detected in the description.")
        with st.container(border=True):
            st.markdown(result.get("diagnosis", "No diagnosis returned."))

    with literature_tab:
        if not sources:
            st.caption("No supporting literature was returned for this query.")
        for source in sources:
            with st.container(border=True):
                st.markdown(f"**{source.get('title', 'Untitled')}**")
                authors = ", ".join(source.get("authors", []))
                st.caption(f"{authors} · {source.get('publication_date', 'n.d.')}")
                with st.expander("Abstract", icon=":material/description:"):
                    st.write(source.get("abstract", "No abstract available."))
                if source.get("article_url"):
                    st.link_button(
                        "View on PubMed",
                        source["article_url"],
                        icon=":material/open_in_new:",
                    )

    with web_tab:
        if not web_sources:
            st.caption("No web results were returned for this query.")
        for source in web_sources:
            with st.container(border=True):
                st.markdown(f"**{source.get('title', 'Untitled')}**")
                st.caption(source.get("content", "No summary available."))
                if source.get("url"):
                    st.link_button(
                        "Open page",
                        source["url"],
                        icon=":material/open_in_new:",
                    )

    with summary_tab:
        st.write(result.get("pubmed_summary", "No summary available."))

    with report_tab:
        report_md = build_report_markdown(entry)
        st.markdown(report_md)

        st.caption("Download this report as:")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.download_button(
                "Markdown",
                data=report_md,
                file_name=f"clininsight_report_{entry['id']}.md",
                mime="text/markdown",
                icon=":material/download:",
                width="stretch",
            )
        with col2:
            st.download_button(
                "Plain text",
                data=build_report_text(entry),
                file_name=f"clininsight_report_{entry['id']}.txt",
                mime="text/plain",
                icon=":material/download:",
                width="stretch",
            )
        with col3:
            st.download_button(
                "Word (.docx)",
                data=build_report_docx(entry),
                file_name=f"clininsight_report_{entry['id']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                width="stretch",
            )
        with col4:
            st.download_button(
                "PDF",
                data=build_report_pdf(entry),
                file_name=f"clininsight_report_{entry['id']}.pdf",
                mime="application/pdf",
                icon=":material/download:",
                width="stretch",
            )


with st.sidebar:
    st.markdown("### ClinInsight AI")
    st.caption("AI-assisted symptom triage & literature research")

    st.text_input("Backend API URL", key="api_base")
    online, status_message = check_backend(st.session_state.api_base)
    with st.container(horizontal=True):
        if online:
            st.badge("Backend online", icon=":material/check_circle:", color="green")
        else:
            st.badge("Backend unreachable", icon=":material/error:", color="red")
    if not online:
        st.caption(f"Start it with: `uvicorn app:app --port 8081` ({status_message})")

    if online:
        st.markdown("#### Local cache")
        stats = get_cache_stats(st.session_state.api_base)
        if stats:
            size_kb = stats.get("size_bytes", 0) / 1024
            st.caption(f"{stats.get('entries', 0)} cached responses · {size_kb:.1f} KB on disk")
        if st.button("Clear server cache", icon=":material/delete_sweep:"):
            try:
                outcome = clear_server_cache(st.session_state.api_base)
                get_cache_stats.clear()
                st.toast(f"Cleared {outcome.get('removed', 0)} cached entries", icon=":material/check_circle:")
            except requests.exceptions.RequestException as exc:
                st.toast(f"Could not clear cache: {exc}", icon=":material/error:")

    st.markdown("#### Recent analyses")
    history = st.session_state.history
    if not history:
        st.caption("Your analyses will appear here.")
    else:
        for entry in reversed(history[-10:]):
            label = entry["description"][:40] + ("…" if len(entry["description"]) > 40 else "")
            if st.button(label or "(empty description)", key=f"hist_{entry['id']}", icon=":material/history:"):
                st.session_state.current_id = entry["id"]
                st.rerun()
        st.button(
            "Clear history",
            icon=":material/delete:",
            on_click=lambda: (st.session_state.history.clear(), st.session_state.update(current_id=None)),
        )

    if history:
        st.markdown("#### Symptom trends")
        counts: dict[str, int] = {}
        for entry in history:
            for symptom in entry["result"].get("symptom", []):
                counts[symptom] = counts.get(symptom, 0) + 1
        if counts:
            trend_df = pd.DataFrame(
                {"symptom": list(counts.keys()), "occurrences": list(counts.values())}
            ).set_index("symptom")
            st.bar_chart(trend_df, width="stretch")

    st.caption("Not a substitute for professional medical advice.")

st.title("ClinInsight AI")
st.caption("Describe symptoms to get an AI-assisted diagnosis overview and supporting PubMed research.")
st.warning(
    "This tool provides informational content only and is not a substitute for professional medical advice.",
    icon=":material/warning:",
)

picked_symptoms = st.pills(
    "Add common symptoms",
    KNOWN_SYMPTOMS,
    format_func=lambda s: f"{SYMPTOM_ICONS.get(s, DEFAULT_SYMPTOM_ICON)} {s}",
    selection_mode="multi",
    key="symptom_pills",
)
# Kept outside st.form so picks are reflected immediately — pills inside a
# form only trigger a rerun on submit, which made it look like extra
# selections were being silently dropped when really the results panel was
# just still showing the previous submission.
if picked_symptoms:
    st.caption(
        "Currently selected: "
        + ", ".join(f"{SYMPTOM_ICONS.get(s, DEFAULT_SYMPTOM_ICON)} {s}" for s in picked_symptoms)
    )

with st.form("symptom_form", border=True):
    description = st.text_area(
        "Describe your symptoms",
        key="description_input",
        placeholder="e.g. I've had a throbbing headache and a mild fever since this morning...",
        height=120,
    )
    severity = st.segmented_control("Severity", ["Mild", "Moderate", "Severe"], key="severity")
    force_refresh = st.checkbox(
        "Force refresh (bypass cache, always regenerate)",
        key="force_refresh",
        help="Cached diagnoses and summaries are reused by default to save on AI/API "
        "calls. Check this to ignore the cache and regenerate from scratch.",
    )
    submitted = st.form_submit_button("Analyze symptoms", icon=":material/send:", type="primary")

if submitted:
    extra_parts = []
    if picked_symptoms:
        extra_parts.append("Additional reported symptoms: " + ", ".join(picked_symptoms) + ".")
    if severity:
        extra_parts.append(f"Overall severity: {severity}.")
    full_description = " ".join([description.strip(), *extra_parts]).strip()

    if not full_description:
        st.error("Please describe your symptoms or select at least one from the list.")
    else:
        with st.status("Running clinical analysis…", expanded=True) as status:
            st.write(":material/manage_search: Extracting symptom keywords")
            st.write(":material/psychology: Consulting the diagnostic model")
            st.write(":material/menu_book: Retrieving related PubMed literature")
            st.write(":material/travel_explore: Searching the web for related resources")
            st.write(":material/summarize: Summarizing findings")
            try:
                result = call_diagnosis_api(
                    st.session_state.api_base, full_description, force_refresh=force_refresh
                )
            except requests.exceptions.RequestException as exc:
                status.update(label="Analysis failed", state="error")
                st.error(f"Could not reach the backend at {st.session_state.api_base}. ({exc})")
                st.stop()
            status.update(label="Analysis complete", state="complete")

        entry_id = uuid.uuid4().hex[:8]
        entry = {
            "id": entry_id,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "description": full_description,
            "severity": severity,
            "result": result,
        }
        st.session_state.history.append(entry)
        st.session_state.current_id = entry_id
        st.toast("Analysis complete", icon=":material/check_circle:")
        st.rerun()

current_entry = next(
    (e for e in st.session_state.history if e["id"] == st.session_state.current_id), None
)
if current_entry:
    render_result(current_entry)
else:
    st.caption("Submit a description above to see the AI-assisted diagnosis, literature and summary.")
